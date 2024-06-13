import os
import datetime as dt
from tkinter import filedialog
from tkinter import messagebox
import tkinter as tk
import docx
from docx2pdf import convert


class InvoiceAutomation:
    def __init__(self):
        # Initialize the main window
        self.root = tk.Tk()
        self.root.title('Invoice Automation')
        self.root.geometry('500x600')

        # Create labels for the form fields
        self.partner_label = tk.Label(self.root, text='Partner')
        self.address_label = tk.Label(self.root, text='Address')
        self.invoice_number_label = tk.Label(self.root, text='Invoice Number')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.unit_label = tk.Label(self.root, text='Unit')
        self.amount_price_label = tk.Label(self.root, text='Amount')
        self.total_price_label = tk.Label(self.root, text='Total Price')

        # Create entry fields for the form
        self.partner_entry = tk.Entry(self.root)
        self.address_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.unit_entry = tk.Entry(self.root)
        self.amount_price_entry = tk.Entry(self.root)
        self.total_price_entry = tk.Entry(self.root)

        # Payment method dictionary
        self.payment_method = {
            'Main Bank': {
                'Recipient': 'Leyomart Ltd.',
                'Bank': 'United Bank For Africa (UBA)',
                'Account Number': '1023050272'
            }
        }

        # Create a dropdown for selecting payment methods
        self.payment_method_var = tk.StringVar(self.root)
        self.payment_method_var.set('Main Bank')
        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method_var, *self.payment_method.keys())

        # Create the "Create Invoice" button
        self.create_button = tk.Button(self.root, text='Create Invoice', command=self.create_invoice)

        # Padding options for layout management
        padding_options = {'fill': 'x', 'expand': True, 'padx': 5, 'pady': 2}

        # Pack the labels and entry fields into the window
        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.address_label.pack(padding_options)
        self.address_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.unit_label.pack(padding_options)
        self.unit_entry.pack(padding_options)
        self.amount_price_label.pack(padding_options)
        self.amount_price_entry.pack(padding_options)
        self.total_price_label.pack(padding_options)
        self.total_price_entry.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        # Start the Tkinter main loop
        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        # Replace old_text with new_text in the given paragraph
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        # Open the template document
        doc = docx.Document('template.docx')

        # Get the selected payment method
        selected_payment_method = self.payment_method[self.payment_method_var.get()]

        try:
            # Create a dictionary of replacements for placeholders in the template
            replacements = {
                "[Date]": dt.datetime.today().strftime('%y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Address]": self.address_entry.get(),
                "[Invoice Number]": self.invoice_number_entry.get(),
                "[Service Description]": self.service_description_entry.get(),
                "[Unit]": self.unit_entry.get(),
                "[Amount]": f"${float(self.amount_price_entry.get()):.2f}",
                "[Total Price]": f"${float(self.total_price_entry.get()):.2f}",
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[Account Number]": selected_payment_method['Account Number']
            }
        except ValueError:
            # Show an error message if there's a problem with the amount
            messagebox.showerror(title='Error', message='Invalid Amount')
            return

        # Replace placeholders in the document's paragraphs
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        # Replace placeholders in the document's tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        # Save the filled document and convert it to PDF
        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF documents', '*.pdf')])
        doc.save('filled.docx')
        convert('filled.docx', save_path)

        # Remove the temporary DOCX file and show a success message
        os.remove('filled.docx')
        messagebox.showinfo('Success', 'Invoice Created and saved successfully')


if __name__ == '__main__':
    # Run the InvoiceAutomation class when the script is executed
    InvoiceAutomation()
